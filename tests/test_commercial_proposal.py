from __future__ import annotations

import hashlib
import zipfile
from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document

from app.commercial_proposal import CommercialProposalService, proposal_payload_from_analysis
from app.excel_estimate import ExcelEstimateRequest


TEMPLATE = Path("data/commercial-proposal-template.docx")


def _estimate() -> ExcelEstimateRequest:
    return ExcelEstimateRequest.model_validate(
        {
            "project_name": "Миграция резервного копирования",
            "project_type_code": "SUP_Complex",
            "vat_rate": 0.22,
            "role_rate_overrides": {"project_manager": (5_800, 4_704)},
            "work_items": [
                {
                    "stage_no": 1,
                    "stage_name": "Обследование",
                    "work_no": "1.1",
                    "work_name": "Сбор требований",
                    "hours_basis": "Всего",
                    "role_assignments": [
                        {"role": "project_manager", "estimated_hours": 16}
                    ],
                }
            ],
            "stage_explanations": [
                {
                    "stage_no": 1,
                    "explanation": "Собираем и подтверждаем требования заказчика.",
                }
            ],
        }
    )


def _payload(summary: str = "Описание проекта."):
    estimate = _estimate()
    return proposal_payload_from_analysis(
        project_name=estimate.project_name,
        proposal_date=date(2026, 8, 27),
        summary=summary,
        rationale="Работы охватывают критичные информационные системы.",
        facts=[{"name": "Количество площадок", "value": "5"}],
        questions=[
            {
                "question": "Какой срок хранения копий?",
                "reason": "Влияет на объём хранилища.",
            }
        ],
        assumptions=["Доступы предоставляет заказчик"],
        warnings=[],
        stage_plan=None,
        estimate=estimate,
    )


def test_builds_proposal_from_retained_template() -> None:
    output = CommercialProposalService(TEMPLATE).build(
        _payload("Миграция системы резервного копирования на новую платформу.")
    )
    document = Document(BytesIO(output))
    body = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "Миграция системы резервного копирования" in body
    assert "Собираем и подтверждаем требования заказчика." in body
    assert "1. ВВЕДЕНИЕ" in body
    assert "2. О НАС" in body
    assert "7. ПРЕИМУЩЕСТВА РАБОТЫ С MONS" in body
    assert "8. РЕФЕРЕНСНЫЕ ПРОЕКТЫ И КОНТАКТЫ" in body
    assert "9. ЗАКЛЮЧЕНИЕ" in body
    assert document.tables[0].cell(1, 1).text == "113 216 руб."
    assert document.tables[0].cell(1, 2).text == "92 800 руб."
    assert document.tables[1].cell(1, 0).text == "Количество площадок"
    with zipfile.ZipFile(BytesIO(output)) as archive:
        assert "МИГРАЦИЯ РЕЗЕРВНОГО КОПИРОВАНИЯ" in archive.read(
            "word/document.xml"
        ).decode("utf-8")


def test_preserves_template_design_parts_byte_for_byte() -> None:
    output = CommercialProposalService(TEMPLATE).build(_payload())
    preserve = {
        "word/styles.xml",
        "word/numbering.xml",
        "word/header1.xml",
        "word/footer1.xml",
        "word/theme/theme1.xml",
        "word/media/image1.png",
        "word/media/image2.png",
        "word/media/image3.png",
    }
    with zipfile.ZipFile(TEMPLATE) as source, zipfile.ZipFile(BytesIO(output)) as result:
        for part in preserve:
            assert hashlib.sha256(source.read(part)).digest() == hashlib.sha256(
                result.read(part)
            ).digest()
