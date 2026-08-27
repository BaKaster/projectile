from __future__ import annotations

import math
import re
import zipfile
from copy import deepcopy
from dataclasses import dataclass, field
from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document as WordDocument
from docx.document import Document as WordDocumentType
from docx.oxml.ns import qn
from docx.table import _Cell, _Row
from docx.text.paragraph import Paragraph
from lxml import etree

from app.excel_estimate import ExcelEstimateRequest


class CommercialProposalError(ValueError):
    """The analysis cannot be safely mapped to the proposal template."""


@dataclass(frozen=True, slots=True)
class ProposalStage:
    name: str
    description: str


@dataclass(frozen=True, slots=True)
class ProposalCostLine:
    name: str
    amount_without_vat: float
    amount_with_vat: float
    timeline: str


@dataclass(frozen=True, slots=True)
class ProposalDetail:
    parameter: str
    description: str
    value: str


@dataclass(slots=True)
class CommercialProposalPayload:
    project_name: str
    proposal_date: date
    project_description: str
    stages: list[ProposalStage]
    cost_lines: list[ProposalCostLine]
    duration_days: int
    obligations: list[str]
    details: list[ProposalDetail] = field(default_factory=list)


_CONSTANT_PARAGRAPHS = (
    "1. ВВЕДЕНИЕ",
    (
        "Настоящий документ является коммерческим предложением компании MONS по "
        "внедрению системы резервного копирования. В документе изложены цели и задачи "
        "проекта, ключевые этапы реализации, условия и стоимость решения. Информация, "
        "представленная в данном предложении, является конфиденциальной и не подлежит "
        "распространению без предварительного согласия MONS."
    ),
    "2. О НАС",
    (
        "MONS — это российская ИТ-компания, часть группы компаний «КОРУС Консалтинг». "
        "Мы создаем и модернизируем ИТ-инфраструктуру, улучшаем ИТ-процессы и "
        "осуществляем техническую поддержку решений. Знаем всё о современных "
        "технологиях и подберем для вас лучшее из возможного для любых задач, даже в "
        "трудные времена. С нами заказчики могут больше времени уделять бизнесу, "
        "потому что ИТ мы берём на себя."
    ),
    "7. ПРЕИМУЩЕСТВА РАБОТЫ С MONS",
    "Экспертиза и опыт: реализовано более 100 успешных проектов по облачной миграции.",
    "Комплексный подход: анализ, проектирование, тестирование, основная миграция, поддержка.",
    "Минимизация рисков: пошаговый план миграции и тестирование перед запуском.",
    "Гибкость решений: адаптация под индивидуальные потребности заказчика.",
    "8. РЕФЕРЕНСНЫЕ ПРОЕКТЫ И КОНТАКТЫ",
    (
        "Мы готовы предоставить рекомендации от наших клиентов по релевантным проектам, "
        "а также, при необходимости, организовать референсные визиты представителей "
        "компании Заказчика на объекты других клиентов «КОРУС Консалтинг»."
    ),
    "Среди наших заказчиков такие компании, как: ",
    "9. ЗАКЛЮЧЕНИЕ",
    (
        "Внедрение системы резервного копирования от MONS обеспечит надёжную защиту и "
        "оперативное восстановление критически важных данных, повысит устойчивость вашей "
        "ИТ-инфраструктуры и снизит эксплуатационные риски. Мы предлагаем комплексный "
        "подход с полным сопровождением на всех этапах реализации проекта."
    ),
    (
        "Если у вас возникнут вопросы или потребуется дополнительная информация, мы "
        "готовы обсудить детали и адаптировать наше предложение под ваши требования."
    ),
    "С уважением,\n[Имя менеджера]\n[Должность]\nКомпания MONS",
)


class CommercialProposalService:
    """Populate project-specific slots in the retained MONS proposal template."""

    def __init__(self, template_path: Path) -> None:
        self.template_path = template_path
        self._template = template_path.read_bytes()
        self._validate_constant_text(WordDocument(BytesIO(self._template)))

    def build(self, payload: CommercialProposalPayload) -> bytes:
        if not payload.project_name.strip():
            raise CommercialProposalError("project name is required")
        if not payload.stages:
            raise CommercialProposalError("at least one project stage is required")
        if not payload.cost_lines:
            raise CommercialProposalError("at least one cost line is required")

        document = WordDocument(BytesIO(self._template))
        self._replace_cover_body_title(document, payload.project_name)
        self._replace_exact_paragraph(
            document,
            "Проект по внедрению системы резервного копирования представляет собой",
            payload.project_description,
        )
        self._replace_stage_block(document, payload.stages)
        self._replace_cost_table(document, payload.cost_lines)
        self._replace_detail_block(document, payload.details)
        self._replace_exact_paragraph(
            document,
            "Общий срок выполнения проекта:",
            (
                f"Общий срок выполнения проекта: {payload.duration_days} рабочих дней "
                "с момента подписания договора и предоставления всех необходимых данных "
                "и доступов."
            ),
        )
        self._replace_exact_paragraph(
            document,
            "Для реализации услуг по технической поддержке заказчик обязуется",
            " ".join(payload.obligations),
        )

        output = BytesIO()
        document.save(output)
        result = self._patch_cover_shapes(
            output.getvalue(), payload.project_name, payload.proposal_date
        )
        self._validate_constant_text(WordDocument(BytesIO(result)))
        return result

    @staticmethod
    def _paragraph_by_prefix(
        document: WordDocumentType, prefix: str
    ) -> Paragraph:
        for paragraph in document.paragraphs:
            if paragraph.text.strip().startswith(prefix):
                return paragraph
        raise CommercialProposalError(f"proposal template slot not found: {prefix}")

    @classmethod
    def _replace_exact_paragraph(
        cls, document: WordDocumentType, prefix: str, text: str
    ) -> None:
        paragraph = cls._paragraph_by_prefix(document, prefix)
        cls._replace_paragraph_text(paragraph, text)

    @staticmethod
    def _replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
        run_properties = next(
            (deepcopy(run._r.rPr) for run in paragraph.runs if run._r.rPr is not None),
            None,
        )
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)
        run = paragraph.add_run(text)
        if run_properties is not None:
            if run._r.rPr is not None:
                run._r.remove(run._r.rPr)
            run._r.insert(0, run_properties)

    @staticmethod
    def _replace_cover_body_title(
        document: WordDocumentType, project_name: str
    ) -> None:
        paragraph = CommercialProposalService._paragraph_by_prefix(
            document, "КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ"
        )
        if len(paragraph.runs) < 3:
            raise CommercialProposalError("proposal title formatting is incompatible")
        paragraph.runs[2].text = f"по проекту «{project_name.strip()}»"

    @staticmethod
    def _replace_stage_block(
        document: WordDocumentType, stages: list[ProposalStage]
    ) -> None:
        heading = CommercialProposalService._paragraph_by_prefix(
            document, "4. ЭТАПЫ РЕАЛИЗАЦИИ ПРОЕКТА"
        )
        intro = CommercialProposalService._paragraph_by_prefix(
            document, "Проект реализуется в следующих основных фазах:"
        )
        cost_heading = CommercialProposalService._paragraph_by_prefix(
            document, "5. СТОИМОСТЬ И УСЛОВИЯ ОПЛАТЫ"
        )
        body = document._element.body
        elements = list(body)
        intro_index = elements.index(intro._p)
        cost_index = elements.index(cost_heading._p)
        templates = [
            item
            for item in elements[intro_index + 1 : cost_index]
            if item.tag == qn("w:p") and item.find(qn("w:pPr")) is not None
            and item.find(f"{qn('w:pPr')}/{qn('w:numPr')}") is not None
        ]
        if not templates:
            raise CommercialProposalError("stage bullet template is missing")
        template = templates[0]
        for item in elements[intro_index + 1 : cost_index]:
            body.remove(item)
        insertion_index = list(body).index(cost_heading._p)
        for offset, stage in enumerate(stages):
            element = deepcopy(template)
            body.insert(insertion_index + offset, element)
            paragraph = Paragraph(element, document._body)
            CommercialProposalService._replace_stage_paragraph(
                paragraph, stage.name, stage.description
            )
        # Keep the visual breathing space between the stage list and section 5.
        spacer = deepcopy(heading._p)
        for child in list(spacer):
            if child.tag != qn("w:pPr"):
                spacer.remove(child)
        body.insert(insertion_index + len(stages), spacer)

    @staticmethod
    def _replace_stage_paragraph(
        paragraph: Paragraph, title: str, description: str
    ) -> None:
        run_properties = [
            deepcopy(run._r.rPr) if run._r.rPr is not None else None
            for run in paragraph.runs[:2]
        ]
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)
        for index, text in enumerate((f"{title.rstrip(':')}:", f"\n{description}")):
            run = paragraph.add_run(text)
            if index < len(run_properties) and run_properties[index] is not None:
                if run._r.rPr is not None:
                    run._r.remove(run._r.rPr)
                run._r.insert(0, run_properties[index])

    @staticmethod
    def _replace_cost_table(
        document: WordDocumentType, lines: list[ProposalCostLine]
    ) -> None:
        table = document.tables[0]
        template_row = deepcopy(table.rows[1]._tr)
        for row in list(table.rows)[1:]:
            table._tbl.remove(row._tr)
        for line in lines:
            table._tbl.append(deepcopy(template_row))
            row = table.rows[-1]
            CommercialProposalService._set_row_values(
                row,
                (
                    line.name,
                    _format_rubles(line.amount_with_vat),
                    _format_rubles(line.amount_without_vat),
                    line.timeline,
                ),
            )

    @staticmethod
    def _replace_detail_block(
        document: WordDocumentType, details: list[ProposalDetail]
    ) -> None:
        heading = CommercialProposalService._paragraph_by_prefix(
            document, "5.2 Детализация стоимости"
        )
        CommercialProposalService._replace_paragraph_text(
            heading, "5.2 Детализация стоимости проекта"
        )
        table = document.tables[1]
        template_row = deepcopy(table.rows[1]._tr)
        for row in list(table.rows)[1:]:
            table._tbl.remove(row._tr)
        rows = details or [
            ProposalDetail(
                "Исходные данные проекта",
                "Ключевые параметры, влияющие на окончательную стоимость и сроки.",
                "Не указано",
            )
        ]
        for detail in rows[:12]:
            table._tbl.append(deepcopy(template_row))
            CommercialProposalService._set_row_values(
                table.rows[-1],
                (detail.parameter, detail.description, detail.value),
            )

    @staticmethod
    def _set_row_values(row: _Row, values: tuple[str, ...]) -> None:
        for cell, value in zip(row.cells, values, strict=True):
            CommercialProposalService._set_cell_text(cell, value)

    @staticmethod
    def _set_cell_text(cell: _Cell, text: str) -> None:
        first = cell.paragraphs[0]
        CommercialProposalService._replace_paragraph_text(first, text)
        for paragraph in list(cell.paragraphs)[1:]:
            cell._tc.remove(paragraph._p)

    def _patch_cover_shapes(
        self, source: bytes, project_name: str, proposal_date: date
    ) -> bytes:
        source_stream = BytesIO(source)
        output = BytesIO()
        preserve = {
            "word/styles.xml",
            "word/numbering.xml",
            "word/header1.xml",
            "word/footer1.xml",
            "word/theme/theme1.xml",
            "word/media/image1.png",
            "word/media/image2.png",
            "word/media/image3.png",
        }
        with (
            zipfile.ZipFile(source_stream, "r") as archive,
            zipfile.ZipFile(BytesIO(self._template), "r") as template,
            zipfile.ZipFile(output, "w") as target,
        ):
            for info in archive.infolist():
                data = (
                    template.read(info.filename)
                    if info.filename in preserve
                    else archive.read(info.filename)
                )
                if info.filename == "word/document.xml":
                    data = self._patch_document_xml(
                        data, project_name, proposal_date
                    )
                target.writestr(info, data)
        return output.getvalue()

    @staticmethod
    def _patch_document_xml(
        source: bytes, project_name: str, proposal_date: date
    ) -> bytes:
        root = etree.fromstring(source)
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        nodes = root.xpath(".//w:t", namespaces=namespace)
        cover_title = _compact_cover_title(project_name)
        for node in nodes:
            if node.text == "РЕЗЕРВНОЕ КОПИРОВАНИЕ":
                node.text = cover_title

        original_date = ["0", "1", ".", "10", ".202", "4"]
        replacement = [
            proposal_date.strftime("%d")[0],
            proposal_date.strftime("%d")[1],
            ".",
            proposal_date.strftime("%m"),
            "." + proposal_date.strftime("%Y")[:3],
            proposal_date.strftime("%Y")[3],
        ]
        values = [node.text or "" for node in nodes]
        for start in range(len(values) - len(original_date) + 1):
            if values[start : start + len(original_date)] != original_date:
                continue
            for node, value in zip(
                nodes[start : start + len(original_date)], replacement, strict=True
            ):
                node.text = value
        return etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    @staticmethod
    def _validate_constant_text(document: WordDocumentType) -> None:
        actual = {paragraph.text for paragraph in document.paragraphs}
        missing = [item for item in _CONSTANT_PARAGRAPHS if item not in actual]
        if missing:
            raise CommercialProposalError(
                "constant proposal sections were changed or are missing: "
                + "; ".join(missing[:3])
            )


def proposal_payload_from_analysis(
    *,
    project_name: str,
    proposal_date: date,
    summary: str,
    rationale: str,
    facts: list[dict],
    questions: list[dict],
    assumptions: list[str],
    warnings: list[str],
    stage_plan: dict | None,
    estimate: ExcelEstimateRequest,
) -> CommercialProposalPayload:
    project_description = " ".join(
        item.strip()
        for item in (summary, rationale)
        if item and item.strip()
    )
    if not project_description:
        project_description = f"Проект «{project_name}» сформирован по предоставленным исходным данным."

    stage_descriptions = {
        item.stage_no: item.explanation for item in estimate.stage_explanations
    }
    work_names: dict[int, list[str]] = {}
    stage_names: dict[int, str] = {}
    for work in estimate.work_items:
        stage_names.setdefault(work.stage_no, work.stage_name)
        work_names.setdefault(work.stage_no, []).append(work.work_name)
    stages = []
    for number in sorted(stage_names):
        description = stage_descriptions.get(number)
        if not description:
            description = "В состав этапа входят работы: " + "; ".join(
                dict.fromkeys(work_names[number])
            ) + "."
        stages.append(ProposalStage(stage_names[number], description))

    cost_lines = _cost_lines(estimate)
    total_hours = sum(
        assignment.estimated_hours
        for work in estimate.work_items
        for assignment in work.role_assignments
    )
    duration_days = max(1, math.ceil(total_hours / estimate.work_hours_per_day))
    details = _proposal_details(facts, questions)

    obligations = [
        "Для реализации проекта Заказчик обязуется обеспечить специалистам MONS доступ к необходимым ресурсам, системам и технической документации.",
        "Заказчик обеспечивает своевременное согласование результатов работ и предоставляет ответы на уточняющие вопросы, влияющие на объём, стоимость и сроки проекта.",
    ]
    unresolved = [item.strip() for item in (*assumptions, *warnings) if item and item.strip()]
    if unresolved:
        obligations.append(
            "До начала соответствующих работ стороны подтверждают следующие допущения и ограничения: "
            + "; ".join(unresolved[:4])
            + "."
        )

    return CommercialProposalPayload(
        project_name=project_name,
        proposal_date=proposal_date,
        project_description=project_description,
        stages=stages,
        cost_lines=cost_lines,
        duration_days=duration_days,
        obligations=obligations,
        details=details,
    )


def _cost_lines(estimate: ExcelEstimateRequest) -> list[ProposalCostLine]:
    rows: list[ProposalCostLine] = []
    default_rates = estimate.role_rate_overrides
    for work in estimate.work_items:
        amount = 0.0
        for assignment in work.role_assignments:
            rate = assignment.sale_rate_override
            if rate is None and assignment.role in default_rates:
                rate = default_rates[assignment.role][0]
            amount += assignment.estimated_hours * float(rate or 0)
        multiplier = work.quantity_override or 1
        if work.hours_basis not in {"На единицу", "Ед. × месяц"}:
            multiplier = 1
        amount *= multiplier
        if work.hours_basis in {"В месяц", "Ед. × месяц"}:
            timeline = "Ежемесячно"
            name = f"{work.work_name} (ежемесячно)"
        else:
            timeline = f"{max(1, math.ceil(sum(a.estimated_hours for a in work.role_assignments) / estimate.work_hours_per_day))} раб. дн."
            name = work.work_name
        amount *= 1 + estimate.commercial_reserve_rate
        amount *= 1 - estimate.discount_rate
        rows.append(
            ProposalCostLine(
                name=name,
                amount_without_vat=round(amount, 2),
                amount_with_vat=round(amount * (1 + estimate.vat_rate), 2),
                timeline=timeline,
            )
        )
    for item in estimate.external_costs:
        unit_price = item.unit_sale_price_override or item.unit_cost
        amount = item.quantity * unit_price
        rows.append(
            ProposalCostLine(
                name=item.description,
                amount_without_vat=round(amount, 2),
                amount_with_vat=round(amount * (1 + estimate.vat_rate), 2),
                timeline="Ежемесячно" if item.periodicity == "В месяц" else "Разово",
            )
        )
    return rows


def _proposal_details(facts: list[dict], questions: list[dict]) -> list[ProposalDetail]:
    result: list[ProposalDetail] = []
    for item in facts:
        name = str(item.get("name") or item.get("label") or "Параметр проекта").strip()
        value = str(item.get("value") or "Не указано").strip()
        source = ", ".join(item.get("source_document_ids") or [])
        result.append(
            ProposalDetail(
                parameter=name,
                description=(f"Подтверждено материалами проекта: {source}" if source else "Подтверждено материалами проекта."),
                value=value,
            )
        )
        if len(result) >= 6:
            break
    for item in questions:
        question = str(item.get("question") or "Требуется уточнение").strip()
        reason = str(item.get("reason") or "Параметр влияет на стоимость или сроки проекта.").strip()
        result.append(ProposalDetail(question, reason, "Не указано"))
        if len(result) >= 12:
            break
    return result


def _compact_cover_title(project_name: str) -> str:
    value = re.sub(r"\s+", " ", project_name).strip().upper()
    if len(value) <= 58:
        return value
    return value[:55].rsplit(" ", 1)[0].rstrip(" ,.;:") + "..."


def _format_rubles(value: float) -> str:
    return f"{value:,.0f}".replace(",", " ") + " руб."
